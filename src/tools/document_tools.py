import os
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import json
from jinja2 import Template, Environment, FileSystemLoader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypdf
import base64
from strands_agents.tools import tool
import structlog

logger = structlog.get_logger()


@tool
def create_pdf_report(
    title: str,
    sections: List[Dict[str, Any]],
    output_path: str,
    metadata: Optional[Dict[str, str]] = None
) -> str:
    """
    Create a PDF report with multiple sections.
    
    Args:
        title: Report title
        sections: List of sections, each with 'title', 'content', and optional 'data'
        output_path: Path to save the PDF
        metadata: Optional metadata (author, date, etc.)
        
    Returns:
        Path to the created PDF file
    """
    try:
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Create document
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata if provided
        if metadata:
            meta_text = "<br/>".join([f"<b>{k}:</b> {v}" for k, v in metadata.items()])
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Add sections
        for section in sections:
            # Section title
            story.append(Paragraph(section['title'], styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Section content
            if 'content' in section:
                for paragraph in section['content'].split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph, styles['Normal']))
                        story.append(Spacer(1, 12))
            
            # Add table if data is provided
            if 'data' in section and isinstance(section['data'], list):
                if section['data'] and isinstance(section['data'][0], dict):
                    # Convert dict data to table
                    headers = list(section['data'][0].keys())
                    table_data = [headers]
                    
                    for row in section['data']:
                        table_data.append([str(row.get(h, '')) for h in headers])
                    
                    # Create table
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    
                    story.append(table)
                    story.append(Spacer(1, 20))
            
            # Add page break after each major section (except last)
            if section != sections[-1]:
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        logger.info("PDF report created", path=output_path, sections=len(sections))
        return output_path
        
    except Exception as e:
        logger.error("Failed to create PDF report", error=str(e))
        raise


@tool
def create_word_document(
    title: str,
    sections: List[Dict[str, Any]],
    output_path: str,
    template_path: Optional[str] = None
) -> str:
    """
    Create a Word document with formatted content.
    
    Args:
        title: Document title
        sections: List of sections with content
        output_path: Path to save the document
        template_path: Optional template document path
        
    Returns:
        Path to the created document
    """
    try:
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Create or load document
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add creation date
        doc.add_paragraph(f"Created: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # Add sections
        for section in sections:
            # Section heading
            doc.add_heading(section['title'], level=1)
            
            # Section content
            if 'content' in section:
                for paragraph in section['content'].split('\n\n'):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add list if provided
            if 'list_items' in section:
                for item in section['list_items']:
                    doc.add_paragraph(item, style='List Bullet')
            
            # Add table if data provided
            if 'data' in section and isinstance(section['data'], list):
                if section['data'] and isinstance(section['data'][0], dict):
                    headers = list(section['data'][0].keys())
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Shading Accent 1'
                    
                    # Add headers
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                    
                    # Add data rows
                    for row_data in section['data']:
                        row_cells = table.add_row().cells
                        for i, header in enumerate(headers):
                            row_cells[i].text = str(row_data.get(header, ''))
            
            doc.add_paragraph()  # Add spacing between sections
        
        # Save document
        doc.save(output_path)
        
        logger.info("Word document created", path=output_path, sections=len(sections))
        return output_path
        
    except Exception as e:
        logger.error("Failed to create Word document", error=str(e))
        raise


@tool
def create_html_report(
    title: str,
    sections: List[Dict[str, Any]],
    output_path: str,
    template: Optional[str] = None,
    style: Optional[Dict[str, Any]] = None
) -> str:
    """
    Create an HTML report with custom styling.
    
    Args:
        title: Report title
        sections: List of sections with content
        output_path: Path to save the HTML file
        template: Optional custom Jinja2 template
        style: Optional custom CSS styles
        
    Returns:
        Path to the created HTML file
    """
    try:
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Default template if not provided
        if not template:
            template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f4f4f4;
        }
        .container {
            background-color: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        h1 {
            color: #1f4788;
            text-align: center;
            margin-bottom: 30px;
        }
        h2 {
            color: #2c5aa0;
            margin-top: 30px;
            border-bottom: 2px solid #e0e0e0;
            padding-bottom: 10px;
        }
        .metadata {
            background-color: #e8f0fe;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #1f4788;
            color: white;
        }
        tr:nth-child(even) {
            background-color: #f2f2f2;
        }
        .chart-container {
            margin: 20px 0;
            text-align: center;
        }
        {% if custom_css %}
        {{ custom_css }}
        {% endif %}
    </style>
</head>
<body>
    <div class="container">
        <h1>{{ title }}</h1>
        
        {% if metadata %}
        <div class="metadata">
            {% for key, value in metadata.items() %}
            <strong>{{ key }}:</strong> {{ value }}<br>
            {% endfor %}
        </div>
        {% endif %}
        
        {% for section in sections %}
        <section>
            <h2>{{ section.title }}</h2>
            
            {% if section.content %}
            <div class="content">
                {{ section.content | replace('\n\n', '</p><p>') | replace('\n', '<br>') | safe }}
            </div>
            {% endif %}
            
            {% if section.data %}
            <table>
                <thead>
                    <tr>
                        {% for header in section.data[0].keys() %}
                        <th>{{ header }}</th>
                        {% endfor %}
                    </tr>
                </thead>
                <tbody>
                    {% for row in section.data %}
                    <tr>
                        {% for value in row.values() %}
                        <td>{{ value }}</td>
                        {% endfor %}
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
            {% endif %}
            
            {% if section.chart %}
            <div class="chart-container">
                <img src="{{ section.chart }}" alt="Chart" style="max-width: 100%;">
            </div>
            {% endif %}
        </section>
        {% endfor %}
    </div>
</body>
</html>
            """
        
        # Prepare context
        context = {
            'title': title,
            'sections': sections,
            'metadata': {
                'Generated': datetime.now().strftime('%B %d, %Y at %I:%M %p'),
                'Report Type': 'Multi-Agent Analysis'
            },
            'custom_css': style.get('css', '') if style else ''
        }
        
        # Render template
        jinja_template = Template(template)
        html_content = jinja_template.render(**context)
        
        # Save HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info("HTML report created", path=output_path, sections=len(sections))
        return output_path
        
    except Exception as e:
        logger.error("Failed to create HTML report", error=str(e))
        raise


@tool
def merge_documents(
    input_paths: List[str],
    output_path: str,
    document_type: str = "pdf"
) -> str:
    """
    Merge multiple documents into a single file.
    
    Args:
        input_paths: List of paths to documents to merge
        output_path: Path for the merged document
        document_type: Type of documents (pdf, docx)
        
    Returns:
        Path to the merged document
    """
    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        if document_type == "pdf":
            merger = pypdf.PdfMerger()
            
            for path in input_paths:
                if os.path.exists(path):
                    merger.append(path)
                else:
                    logger.warning(f"File not found: {path}")
            
            merger.write(output_path)
            merger.close()
            
        elif document_type == "docx":
            # For Word documents, append content
            merged_doc = Document()
            
            for i, path in enumerate(input_paths):
                if os.path.exists(path):
                    doc = Document(path)
                    
                    # Add page break between documents (except for first)
                    if i > 0:
                        merged_doc.add_page_break()
                    
                    # Copy all elements
                    for element in doc.element.body:
                        merged_doc.element.body.append(element)
            
            merged_doc.save(output_path)
        
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
        
        logger.info(
            "Documents merged",
            input_count=len(input_paths),
            output_path=output_path,
            type=document_type
        )
        
        return output_path
        
    except Exception as e:
        logger.error("Failed to merge documents", error=str(e))
        raise


@tool
def extract_text_from_pdf(pdf_path: str, page_range: Optional[tuple] = None) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        page_range: Optional tuple of (start_page, end_page)
        
    Returns:
        Extracted text content
    """
    try:
        text_content = []
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Determine pages to extract
            if page_range:
                start_page = max(0, page_range[0] - 1)
                end_page = min(len(pdf_reader.pages), page_range[1])
            else:
                start_page = 0
                end_page = len(pdf_reader.pages)
            
            # Extract text from each page
            for page_num in range(start_page, end_page):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        
        full_text = "\n\n".join(text_content)
        
        logger.info(
            "PDF text extracted",
            path=pdf_path,
            pages_extracted=end_page - start_page,
            total_length=len(full_text)
        )
        
        return full_text
        
    except Exception as e:
        logger.error("Failed to extract PDF text", path=pdf_path, error=str(e))
        raise


@tool
def create_template_document(
    template_type: str,
    variables: Dict[str, Any],
    output_path: str
) -> str:
    """
    Create a document from a predefined template.
    
    Args:
        template_type: Type of template (executive_summary, technical_report, etc.)
        variables: Variables to fill in the template
        output_path: Path to save the generated document
        
    Returns:
        Path to the created document
    """
    try:
        templates = {
            "executive_summary": {
                "title": "Executive Summary: {project_name}",
                "sections": [
                    {
                        "title": "Overview",
                        "content": "{overview}"
                    },
                    {
                        "title": "Key Findings",
                        "content": "{key_findings}"
                    },
                    {
                        "title": "Recommendations",
                        "content": "{recommendations}"
                    },
                    {
                        "title": "Next Steps",
                        "content": "{next_steps}"
                    }
                ]
            },
            "technical_report": {
                "title": "Technical Analysis: {subject}",
                "sections": [
                    {
                        "title": "Introduction",
                        "content": "{introduction}"
                    },
                    {
                        "title": "Methodology",
                        "content": "{methodology}"
                    },
                    {
                        "title": "Results",
                        "content": "{results}"
                    },
                    {
                        "title": "Discussion",
                        "content": "{discussion}"
                    },
                    {
                        "title": "Conclusion",
                        "content": "{conclusion}"
                    }
                ]
            }
        }
        
        if template_type not in templates:
            raise ValueError(f"Unknown template type: {template_type}")
        
        # Get template and fill variables
        template = templates[template_type]
        
        # Process title
        title = template["title"].format(**variables)
        
        # Process sections
        sections = []
        for section in template["sections"]:
            processed_section = {
                "title": section["title"],
                "content": section["content"].format(**variables)
            }
            sections.append(processed_section)
        
        # Determine output format from file extension
        ext = Path(output_path).suffix.lower()
        
        if ext == ".pdf":
            return create_pdf_report(title, sections, output_path)
        elif ext == ".docx":
            return create_word_document(title, sections, output_path)
        elif ext == ".html":
            return create_html_report(title, sections, output_path)
        else:
            raise ValueError(f"Unsupported output format: {ext}")
        
    except Exception as e:
        logger.error("Failed to create template document", error=str(e))
        raise